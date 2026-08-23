import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()

    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Title Banner
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("Google A2UI (Agent-to-User Interface) Master Workshop Guide")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0xDA, 0x29, 0x1C) # McDonald's Red

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(16)
    run_sub = sub_p.add_run("A 35-Minute Hands-On Blueprint: Architecture, Agentic Backend & @a2ui/react Frontend")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Metadata Box Table
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False

    col_widths = [Inches(3.3), Inches(3.3)]
    for row in meta_table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]
            set_cell_background(cell, "FFF8E7")
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)

    meta_table.rows[0].cells[0].paragraphs[0].add_run("🎯 Target Audience: Developers, Architects, AI Engineers").bold = True
    meta_table.rows[0].cells[1].paragraphs[0].add_run("⏱️ Session Duration: 30 - 40 Minutes").bold = True
    meta_table.rows[1].cells[0].paragraphs[0].add_run("🛠️ Stack: Python FastAPI + Google ADK + React (@a2ui/react)").italic = True
    meta_table.rows[1].cells[1].paragraphs[0].add_run("🍔 Use Case: McDonald's AI Food Ordering Kiosk").italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Heading 1: Agenda
    h1 = doc.add_heading("1. Workshop Agenda & Session Timing (35 Mins Total)", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0xDA, 0x29, 0x1C)

    agenda_table = doc.add_table(rows=5, cols=3)
    agenda_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    agenda_widths = [Inches(1.5), Inches(3.5), Inches(1.8)]

    headers = ["Timing", "Topic / Module", "Format"]
    for i, title in enumerate(headers):
        cell = agenda_table.rows[0].cells[i]
        cell.paragraphs[0].add_run(title).bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, "DA291C")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)

    agenda_data = [
        ("00:00 - 08:00 (8m)", "Part 1: What is A2UI? Real-World Problem, Architecture & 3 Core Messages", "Conceptual Slides"),
        ("08:00 - 20:00 (12m)", "Part 2: Agentic Backend Walkthrough (Orchestrator, Sub-Agents, Tools & JSON DB)", "Code Walkthrough & CLI"),
        ("20:00 - 30:00 (10m)", "Part 3: Frontend Walkthrough (@a2ui/react, McDonaldsRenderer, Action Loop)", "Code Walkthrough & UI"),
        ("30:00 - 35:00 (5m)", "Part 4: Live End-to-End Demo (Menu -> Customizer -> Tray & Multi-Card Pay -> Invoice)", "Interactive Live Demo"),
    ]

    for row_idx, data in enumerate(agenda_data, start=1):
        for col_idx, text in enumerate(data):
            cell = agenda_table.rows[row_idx].cells[col_idx]
            cell.paragraphs[0].add_run(text)
            set_cell_background(cell, "F9F9F9" if row_idx % 2 == 0 else "FFFFFF")
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Heading 2: Concept
    h2 = doc.add_heading("2. Part 1: What is A2UI and Why Does It Matter?", level=1)
    h2.runs[0].font.color.rgb = RGBColor(0xDA, 0x29, 0x1C)

    p = doc.add_paragraph()
    p.add_run("The Fundamental Problem with Traditional Chatbots:\n").bold = True
    p.add_run(
        "When a user orders food through a conventional text chatbot, the AI outputs paragraphs of text asking for options: "
        "'Which burger would you like? Type 1 for medium fries, type 2 for large...'. This text-based interface is high-friction, error-prone, "
        "and slow. Real humans need visual menus, touch radio buttons, ingredient checkboxes, and one-tap payment."
    )

    doc.add_paragraph()
    p_comp = doc.add_paragraph()
    p_comp.add_run("Traditional LLM Chatbot vs. Google A2UI Generative UI:").bold = True

    comp_table = doc.add_table(rows=5, cols=2)
    comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    comp_headers = ["Traditional Text Chatbot ❌", "Google A2UI Generative UI ✨"]
    for i, title in enumerate(comp_headers):
        cell = comp_table.rows[0].cells[i]
        cell.paragraphs[0].add_run(title).bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, "27251F")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

    comp_data = [
        ("Wall of text messages; tedious to read and reply", "Native, high-resolution interactive UI widgets and cards"),
        ("High latency: Every minor option change requires an LLM call", "60 FPS local reactive state via JSON Pointer data binding"),
        ("Unstructured user text creates parsing failures", "Deterministic client action events ({ eventName, context })"),
        ("Vulnerable to XSS / arbitrary HTML code injection", "Zero-XSS Sandbox: Pure declarative JSON catalog schema"),
    ]

    for row_idx, data in enumerate(comp_data, start=1):
        for col_idx, text in enumerate(data):
            cell = comp_table.rows[row_idx].cells[col_idx]
            cell.paragraphs[0].add_run(text)
            set_cell_background(cell, "F5F5F7" if row_idx % 2 == 0 else "FFFFFF")
            set_cell_margins(cell, top=70, bottom=70, left=100, right=100)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 3 Core Messages
    h2_sub = doc.add_heading("The 3 Core Protocol Messages in A2UI v0.9:", level=2)
    h2_sub.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    msg_bullets = [
        ("1. createSurface: ", "Initializes an isolated surface container on the client, sets the surfaceId, specifies the component catalog (e.g. basicCatalog), and defines styling themes."),
        ("2. updateComponents: ", "Declares a hierarchical component layout using a flat array of component objects with IDs (e.g., Column -> Row -> Card -> Image, Text, ChoicePicker, Button)."),
        ("3. updateDataModel: ", "Populates and updates dynamic reactive data using JSON Pointer paths (e.g. path: '/meals' or path: '/custom/piriPiri').")
    ]
    for b_title, b_desc in msg_bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.add_run(b_title).bold = True
        bp.add_run(b_desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Heading 3: Backend
    h3 = doc.add_heading("3. Part 2: Agentic Backend Architecture & Code Walkthrough", level=1)
    h3.runs[0].font.color.rgb = RGBColor(0xDA, 0x29, 0x1C)

    doc.add_paragraph(
        "The backend is built with FastAPI and Google ADK, implementing the Router-Specialist Agentic Design Pattern. "
        "Business logic is cleanly decoupled into deterministic Python tools and a structured JSON menu database."
    )

    doc.add_heading("Architecture Overview:", level=2)
    arch_box = doc.add_paragraph()
    arch_box.paragraph_format.left_indent = Inches(0.2)
    arch_box.add_run(
        "Master Kiosk Orchestrator (server/agent.py)\n"
        "  ├── 1. MenuDiscoveryAgent       --> tool: query_menu_database()\n"
        "  ├── 2. MealCustomizerAgent      --> tool: get_meal_details(), get_customization_options()\n"
        "  ├── 3. CartAndPaymentAgent      --> tool: calculate_tray_totals(), get_available_payment_methods()\n"
        "  └── 4. InvoiceSettlementAgent   --> tool: generate_purchase_order_invoice()\n"
    ).font.name = 'Courier New'

    doc.add_heading("Key Backend Files & Components:", level=2)
    files_data = [
        ("server/restaurant_data.json", "Source of truth for Indian McDonald's meals, dietary flags (Veg/Non-Veg), prices (₹), calories, customization add-ons, and payment options."),
        ("server/tools.py", "Modular Python tools handling multi-field search, 5% GST tax calculation (2.5% CGST + 2.5% SGST), and PO generation."),
        ("server/agent.py", "Implements ADKRestaurantAgent with the 4 specialist sub-agents, routing natural language queries and client UI action events."),
        ("server/main.py", "FastAPI web server exposing GET / (health & spec) and POST /agent/message (A2UI payload streaming).")
    ]
    for fn, fdesc in files_data:
        p_f = doc.add_paragraph(style='List Bullet')
        p_f.add_run(fn + ": ").bold = True
        p_f.add_run(fdesc)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Heading 4: Frontend
    h4 = doc.add_heading("4. Part 3: Frontend Walkthrough with @a2ui/react", level=1)
    h4.runs[0].font.color.rgb = RGBColor(0xDA, 0x29, 0x1C)

    doc.add_paragraph(
        "The frontend is a pure React + TypeScript application leveraging Google's official @a2ui/react and @a2ui/web_core SDKs. "
        "It features an inline conversational chatbot with a persistent side-by-side Live A2UI Protocol Inspector."
    )

    doc.add_heading("McDonaldsRenderer Component (client/src/a2ui/mcdonaldsRenderer.tsx):", level=2)
    doc.add_paragraph(
        "The McDonaldsRenderer initializes an instance of MessageProcessor with basicCatalog, feeds incoming A2UI messages into the model, "
        "and renders safe native React components via <A2uiSurface />:"
    )

    code_p = doc.add_paragraph()
    code_p.paragraph_format.left_indent = Inches(0.2)
    code_p.add_run(
        "const newProcessor = new MessageProcessor([basicCatalog], async (action: A2uiClientAction) => {\n"
        "  if (onAction) onAction(action); // Dispatches event back to agent loop\n"
        "});\n"
        "newProcessor.processMessages(messages);\n"
        "return surfaces.map(surface => <A2uiSurface key={surface.id} surface={surface} />);\n"
    ).font.name = 'Courier New'

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Heading 5: Live Demo Script
    h5 = doc.add_heading("5. Part 4: Step-by-Step Live Demo Presentation Script", level=1)
    h5.runs[0].font.color.rgb = RGBColor(0xDA, 0x29, 0x1C)

    demo_steps = [
        ("Step 1: Browse Indian Best Sellers Menu",
         "Type: 'Show me the menu'",
         "Show how the agent emits createSurface, updateComponents, and updateDataModel. Point to the live JSON stream in the Inspector."),

        ("Step 2: Interactive Meal Customization & Two-Way Binding",
         "Click: 'Customize Meal 🍔' on McSpicy Paneer Card",
         "Toggle Piri Piri Spice Mix (+₹25) and Extra Cheese (+₹35). Demonstrate that checkbox state updates instantly at 60 FPS without LLM latency."),

        ("Step 3: Tray Review with Multiple Card Payment Options",
         "Click: 'Add Customized Meal to Order 🛒'",
         "Demonstrate itemized pricing, 5% GST breakdown, and the Multi-Card ChoicePicker (Visa/Mastercard, RuPay, Amex, Google Pay UPI, Apple Pay)."),

        ("Step 4: Payment Confirmation & GST Tax Invoice #88",
         "Click: '💳 Pay with Selected Card/UPI & Print Invoice'",
         "Watch the confetti shoot across the screen. Show the official GST Tax Invoice and Order Pickup Token #88 with Table Tent #12 instructions.")
    ]

    for step_num, step_act, step_talk in demo_steps:
        sp = doc.add_paragraph()
        sp.add_run(step_num + "\n").bold = True
        sp.add_run(f"• Action: {step_act}\n").italic = True
        sp.add_run(f"• Presenter Talking Point: {step_talk}")

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # Heading 6: Key Takeaways
    h6 = doc.add_heading("6. Key Architecture Takeaways for Attendees", level=1)
    h6.runs[0].font.color.rgb = RGBColor(0xDA, 0x29, 0x1C)

    takeaways = [
        ("Security by Design: ", "Zero eval() or raw HTML injection. The client only renders trusted catalog components."),
        ("Multi-Platform Portability: ", "The exact same A2UI JSON payload can be consumed by Web (React), Mobile (Flutter, SwiftUI, Jetpack Compose), and in-car kiosks."),
        ("Deterministic Two-Way Event Loop: ", "Client components dispatch typed JSON payloads ({ eventName, context }), allowing agents to update state with surgical precision.")
    ]
    for t_title, t_desc in takeaways:
        tp = doc.add_paragraph(style='List Bullet')
        tp.add_run(t_title).bold = True
        tp.add_run(t_desc)

    # Save
    out_path = "/Users/arjunvijay/Documents/GitHub/gdg-a2ui-demo/Google_A2UI_Workshop_Guide.docx"
    doc.save(out_path)
    print(f"Successfully generated Word document at: {out_path}")

if __name__ == "__main__":
    create_document()
